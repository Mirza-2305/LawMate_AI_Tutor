# app.py - Final Working Version
import streamlit as st
from pathlib import Path
import sys
import os
import uuid
from datetime import datetime
from io import BytesIO

# Add path for local imports
sys.path.append(str(Path(__file__).parent))

from supabase_client import SupabaseManager
from text_extraction import extract_text, get_preview_text
from chunking import chunk_text, find_relevant_chunks
from qa import get_answer_from_chunks, get_available_models, DEFAULT_MODEL
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document

# === EXPORT FUNCTIONS ===
def create_pdf_export(data: dict) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=18, spaceAfter=20)
    story.append(Paragraph("Document Q&A Response", title_style))
    story.append(Spacer(1, 12))

    normal_style = ParagraphStyle('Normal', parent=styles['Normal'], fontSize=10, spaceBefore=6, spaceAfter=6)
    story.append(Paragraph(f"Date: {data['timestamp']}", normal_style))
    story.append(Paragraph(f"Model: {data['model']}", normal_style))

    if data.get('used_general_knowledge', False):
        response_type = "General Knowledge + Document Context"
    elif not data.get('has_document_context', True):
        response_type = "General Knowledge"
    else:
        response_type = "Document-Based"

    story.append(Paragraph(f"Type: {response_type}", normal_style))
    story.append(Spacer(1, 20))

    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=12, spaceBefore=15, spaceAfter=8)
    story.append(Paragraph("Question:", heading_style))
    story.append(Paragraph(data['question'], normal_style))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Answer:", heading_style))
    for line in data['answer'].split('\n'):
        if line.strip():
            story.append(Paragraph(line.strip(), normal_style))
        else:
            story.append(Spacer(1, 6))

    if data['sources']:
        story.append(Spacer(1, 12))
        story.append(Paragraph("Sources:", heading_style))
        for source in data['sources']:
            story.append(Paragraph(f"• {source}", normal_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def create_docx_export(data: dict) -> bytes:
    buffer = BytesIO()
    doc = Document()

    doc.add_heading('Document Q&A Response', 0).alignment = 1
    doc.add_paragraph(f'Date: {data["timestamp"]}')
    doc.add_paragraph(f'Model: {data["model"]}')

    if data.get('used_general_knowledge', False):
        response_type = "General Knowledge + Document Context"
    elif not data.get('has_document_context', True):
        response_type = "General Knowledge"
    else:
        response_type = "Document-Based"

    doc.add_paragraph(f'Type: {response_type}')
    doc.add_paragraph()

    doc.add_heading('Question:', level=1)
    doc.add_paragraph(data['question'])
    doc.add_heading('Answer:', level=1)
    for line in data['answer'].split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())

    if data['sources']:
        doc.add_heading('Sources:', level=1)
        for source in data['sources']:
            doc.add_paragraph(f'• {source}', style='List Bullet')

    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

# === AUTHENTICATION ===
def login_user():
    st.sidebar.title("🔐 User Login")
    if 'user' not in st.session_state:
        st.session_state.user = None

    if st.session_state.user is None:
        username = st.sidebar.text_input("Username")
        password = st.sidebar.text_input("Password", type="password")

        if st.sidebar.button("Login"):
            file_manager = st.session_state.file_manager
            user = file_manager.verify_user(username, password)
            if user:
                st.session_state.user = user
                st.sidebar.success(f"✅ Welcome, {user['username']} ({user['role']})!")
                st.rerun()
            else:
                st.sidebar.error("❌ Invalid credentials")

        st.sidebar.info("Default admin login: username='admin', password='########'")
    else:
        st.sidebar.success(f"Logged in: {st.session_state.user['username']} ({st.session_state.user['role']})")
        if st.sidebar.button("Logout"):
            st.session_state.user = None
            st.rerun()

# === MAIN FUNCTION ===
def main():
    st.set_page_config(
        page_title="LawMate AI Tutor",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # Optional font styling
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Source+Sans+Pro:wght@400;600;700&display=swap');
    </style>
    """, unsafe_allow_html=True)

    # Initialize session state
    if 'file_manager' not in st.session_state:
        try:
            st.session_state.file_manager = SupabaseManager()
        except Exception as e:
            st.error(f"❌ Failed to initialize Supabase: {e}")
            st.stop()
            
    if 'qa_history' not in st.session_state:
        st.session_state.qa_history = []

    # Login system
    login_user()

    # Determine user context
    is_admin = False
    user_id = "guest"
    user_role = "guest"
    if st.session_state.user:
        is_admin = st.session_state.user['role'] == 'admin'
        user_id = st.session_state.user['user_id']
        user_role = st.session_state.user['role']

    # Admin/User/Guest info
    if is_admin:
        st.success("👑 Admin Mode")
    elif st.session_state.user:
        st.info(f"👤 User Mode - You see your docs + admin docs")
    else:
        st.warning("👥 Guest Mode - Only admin docs visible")

    # === UPLOAD SECTION ===
    st.sidebar.title("📤 Upload Document")
    uploaded_file = st.sidebar.file_uploader("Choose a file", type=['pdf','docx','txt','png','jpg','jpeg'])
    countries = ["Pakistan", "India", "USA", "UK", "Canada", "UAE"]
    country = st.sidebar.selectbox("Select Country", countries)
    doc_types = ["Ordinance","Act","Course Material","Student Material","Policy","Report","Other"]
    doc_type = st.sidebar.selectbox("Document Type", doc_types)

    if st.sidebar.button("Upload Document", type="primary"):
        if uploaded_file is not None:
            try:
                file_manager = st.session_state.file_manager
                file_content = uploaded_file.getvalue()
                if not file_content:
                    st.sidebar.error("❌ File is empty")
                    st.stop()

                # Extract text
                with st.spinner("Extracting text..."):
                    text = extract_text(file_content, Path(uploaded_file.name).suffix)
                
                if not text or len(text.strip()) < 50:
                    st.sidebar.error("❌ Failed to extract meaningful text from file")
                else:
                    # Chunk text
                    with st.spinner("Processing document..."):
                        chunks = chunk_text(text, str(uuid.uuid4()), 
                                          chunk_size=800, overlap=100)
                    
                    # Upload to Supabase
                    owner_role = "admin" if is_admin else "user"
                    doc_id = file_manager.add_document(
                        uploaded_file.name, country, doc_type,
                        user_id, owner_role, file_content, chunks
                    )
                    
                    if doc_id:
                        st.sidebar.success(f"✅ Uploaded! ID: {doc_id[:8]}...")
                        st.rerun()
                    else:
                        st.sidebar.error("❌ Upload failed - check error details above")
                        
            except Exception as e:
                st.sidebar.error(f"❌ Upload error: {e}")
        else:
            st.sidebar.warning("Please select a file first")

    # === DOCUMENT LIST + SEARCH ===
    st.title("📄 Uploaded Documents")
    file_manager = st.session_state.file_manager
    search_keyword = st.text_input("Search Documents", placeholder="Enter keyword...")
    filter_country = st.selectbox("Filter by Country", ["All"] + countries)
    filter_type = st.selectbox("Filter by Type", ["All"] + doc_types)

    try:
        if search_keyword:
            documents = file_manager.search_documents(user_id, user_role, search_keyword)
        else:
            documents = file_manager.get_documents_by_filters(
                user_id, user_role, filter_country, filter_type
            )
    except Exception as e:
        st.error(f"❌ Error fetching documents: {e}")
        documents = []

    if not documents:
        st.info("No documents found. Upload some documents to get started!")
    else:
        for idx, doc in enumerate(documents):
            with st.expander(
                f"{doc['filename']} ({doc['country']} - {doc['doc_type']})"
            ):
                # Show preview
                preview_text = doc.get('chunks', [{}])[0].get('text', '') if doc.get('chunks') else ''
                preview = get_preview_text(preview_text, 300)
                st.text_area("Preview", preview, height=100, key=f"preview_{idx}")
                
                # Show metadata
                st.caption(
                    f"ID: {doc['id'][:12]}... | Owner: {doc['owner_role']} | "
                    f"Date: {doc.get('upload_date', 'N/A')[:10]}"
                )

    # === Q&A SECTION ===
    st.markdown("---")
    st.markdown("### 🙋 Ask a Question")
    
    available_models = get_available_models()
    selected_model = st.selectbox("Select AI Model", available_models, index=0)
    question = st.text_area("Enter your question:", height=100)

    if st.button("Get Answer", type="primary"):
        if not question.strip():
            st.warning("Please enter a question first")
        else:
            try:
                with st.spinner("Searching documents and generating answer..."):
                    all_chunks = file_manager.get_all_chunks(user_id, user_role)
                    
                    if not all_chunks:
                        st.warning("⚠️ No document chunks available. Upload documents first.")
                        st.stop()
                    
                    relevant_chunks = find_relevant_chunks(question, all_chunks, top_k=5)
                    
                    if not relevant_chunks:
                        st.info("ℹ️ No relevant document sections found. Will use general knowledge.")
                    
                    result = get_answer_from_chunks(
                        query=question, chunks=relevant_chunks, model=selected_model
                    )

                st.markdown("#### 🤖 AI Answer")
                st.markdown(result['answer'])
                
                # Store in history
                st.session_state.qa_history.append({
                    "question": question,
                    "answer": result['answer'],
                    "model": selected_model,
                    "timestamp": datetime.now().isoformat()
                })

            except Exception as e:
                st.error(f"❌ Error generating answer: {e}")
                st.info("💡 Tip: Check that documents are properly uploaded and chunked")

